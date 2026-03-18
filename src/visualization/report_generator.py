"""
报告生成器 - 生成综合分析报告

功能：
1. 多图表布局报告
2. 数据分析摘要
3. 多格式导出 (PDF, HTML, Markdown, DOCX)
4. 批量报告生成
5. 报告模板系统
"""

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

from .base import BaseVisualizer
from .candlestick_chart import CandlestickChart
from .key_levels_plot import KeyLevelsPlotter
from .market_structure_plot import MarketStructurePlotter


class ReportGenerator(BaseVisualizer):
    """报告生成器"""
    
    def __init__(self, 
                 figsize: Tuple[int, int] = (16, 9),
                 dpi: int = 100,
                 style: str = 'seaborn-v0_8-darkgrid',
                 color_palette: str = 'Set2'):
        """
        初始化报告生成器
        
        Args:
            figsize: 图表尺寸
            dpi: 图表分辨率
            style: matplotlib样式
            color_palette: 颜色调色板
        """
        super().__init__(figsize, dpi, style, color_palette)
        
        # 报告配置
        self.report_config = {
            'title': '价格行为学分析报告',
            'author': '价格行为学量化分析系统',
            'date_format': '%Y-%m-%d %H:%M:%S',
            'default_format': 'pdf',
            'templates': {
                'standard': self.standard_template,
                'detailed': self.detailed_template,
                'summary': self.summary_template
            }
        }
        
        # 初始化子组件
        self.candlestick_chart = CandlestickChart(figsize, dpi, style, color_palette)
        self.key_levels_plotter = KeyLevelsPlotter(figsize, dpi, style, color_palette)
        self.market_structure_plotter = MarketStructurePlotter(figsize, dpi, style, color_palette)
        
        # 报告数据存储
        self.report_data = {}
        self.report_charts = {}
        
    def generate_report(self,
                       data: pd.DataFrame,
                       symbol: str,
                       analysis_results: Dict,
                       template: str = 'standard',
                       output_format: str = 'pdf',
                       output_path: Optional[str] = None) -> Dict:
        """
        生成综合分析报告
        
        Args:
            data: 价格数据
            symbol: 标的符号
            analysis_results: 分析结果
            template: 报告模板
            output_format: 输出格式
            output_path: 输出路径
            
        Returns:
            Dict: 报告生成信息
        """
        # 存储数据
        self.report_data = {
            'symbol': symbol,
            'data': data,
            'analysis_results': analysis_results,
            'generation_time': datetime.now(),
            'template': template,
            'output_format': output_format
        }
        
        # 选择模板
        if template in self.report_config['templates']:
            template_func = self.report_config['templates'][template]
            report_content = template_func(data, symbol, analysis_results)
        else:
            # 使用标准模板
            report_content = self.standard_template(data, symbol, analysis_results)
        
        # 生成图表
        charts = self.generate_report_charts(data, analysis_results)
        self.report_charts = charts
        
        # 生成报告
        report_result = self.compile_report(report_content, charts, output_format, output_path)
        
        return report_result
    
    def standard_template(self,
                         data: pd.DataFrame,
                         symbol: str,
                         analysis_results: Dict) -> Dict:
        """
        标准报告模板
        
        Args:
            data: 价格数据
            symbol: 标的符号
            analysis_results: 分析结果
            
        Returns:
            Dict: 报告内容
        """
        # 提取分析结果
        market_structure = analysis_results.get('market_structure', {})
        key_levels = analysis_results.get('key_levels', {})
        
        # 生成报告内容
        report_content = {
            'title': f'{symbol} 价格行为学分析报告',
            'subtitle': '标准分析报告',
            'sections': [
                {
                    'title': '执行摘要',
                    'content': self.generate_executive_summary(data, symbol, analysis_results)
                },
                {
                    'title': '市场结构分析',
                    'content': self.generate_market_structure_summary(market_structure)
                },
                {
                    'title': '关键价位分析',
                    'content': self.generate_key_levels_summary(key_levels)
                },
                {
                    'title': '技术分析',
                    'content': self.generate_technical_analysis(data)
                },
                {
                    'title': '交易建议',
                    'content': self.generate_trading_recommendations(analysis_results)
                }
            ],
            'metadata': {
                'symbol': symbol,
                'period': f"{data.index[0].strftime('%Y-%m-%d')} 至 {data.index[-1].strftime('%Y-%m-%d')}",
                'data_points': len(data),
                'generation_time': datetime.now().strftime(self.report_config['date_format'])
            }
        }
        
        return report_content
    
    def detailed_template(self,
                         data: pd.DataFrame,
                         symbol: str,
                         analysis_results: Dict) -> Dict:
        """
        详细报告模板
        
        Args:
            data: 价格数据
            symbol: 标的符号
            analysis_results: 分析结果
            
        Returns:
            Dict: 报告内容
        """
        # 基于标准模板，添加更多细节
        standard_content = self.standard_template(data, symbol, analysis_results)
        
        # 添加额外章节
        detailed_sections = standard_content['sections']
        
        # 添加风险分析章节
        detailed_sections.append({
            'title': '风险分析',
            'content': self.generate_risk_analysis(data, analysis_results)
        })
        
        # 添加性能指标章节
        detailed_sections.append({
            'title': '性能指标',
            'content': self.generate_performance_metrics(data)
        })
        
        standard_content['sections'] = detailed_sections
        standard_content['subtitle'] = '详细分析报告'
        
        return standard_content
    
    def summary_template(self,
                        data: pd.DataFrame,
                        symbol: str,
                        analysis_results: Dict) -> Dict:
        """
        摘要报告模板
        
        Args:
            data: 价格数据
            symbol: 标的符号
            analysis_results: 分析结果
            
        Returns:
            Dict: 报告内容
        """
        report_content = {
            'title': f'{symbol} 价格行为学分析摘要',
            'subtitle': '快速分析摘要',
            'sections': [
                {
                    'title': '关键发现',
                    'content': self.generate_key_findings(data, symbol, analysis_results)
                },
                {
                    'title': '交易信号',
                    'content': self.generate_trading_signals(analysis_results)
                }
            ],
            'metadata': {
                'symbol': symbol,
                'period': f"{data.index[0].strftime('%Y-%m-%d')} 至 {data.index[-1].strftime('%Y-%m-%d')}",
                'generation_time': datetime.now().strftime(self.report_config['date_format'])
            }
        }
        
        return report_content
    
    def generate_report_charts(self,
                              data: pd.DataFrame,
                              analysis_results: Dict) -> Dict:
        """
        生成报告图表
        
        Args:
            data: 价格数据
            analysis_results: 分析结果
            
        Returns:
            Dict: 图表数据
        """
        charts = {}
        
        try:
            # 1. 主K线图
            fig1, axes1 = self.candlestick_chart.plot_candlestick(
                data,
                title='价格图表',
                volume=True,
                type='candle',
                indicators=['sma', 'ema', 'macd', 'rsi'],
                show=False
            )
            charts['main_chart'] = {'fig': fig1, 'axes': axes1}
            
            # 2. 关键价位图表
            if 'key_levels' in analysis_results:
                key_levels_data = analysis_results['key_levels']
                support_levels = key_levels_data.get('support_resistance', {}).get('support_levels', [])
                resistance_levels = key_levels_data.get('support_resistance', {}).get('resistance_levels', [])
                pivot_points = key_levels_data.get('pivot_points', {})
                
                fig2, ax2 = self.key_levels_plotter.create_figure()
                ax2.plot(data.index, data['Close'], color='blue', alpha=0.7, label='收盘价')
                
                self.key_levels_plotter.plot_key_levels(
                    ax2, data,
                    support_levels, resistance_levels,
                    pivot_points,
                    show_labels=True
                )
                
                ax2.set_title('关键价位分析')
                ax2.legend()
                charts['key_levels_chart'] = {'fig': fig2, 'ax': ax2}
            
            # 3. 市场结构图表
            if 'market_structure' in analysis_results:
                fig3, ax3 = self.market_structure_plotter.create_figure()
                ax3.plot(data.index, data['Close'], color='blue', alpha=0.7, label='收盘价')
                
                self.market_structure_plotter.plot_market_structure(
                    ax3, data, analysis_results['market_structure'],
                    show_labels=True
                )
                
                ax3.set_title('市场结构分析')
                ax3.legend()
                charts['market_structure_chart'] = {'fig': fig3, 'ax': ax3}
            
        except Exception as e:
            print(f"生成图表时出错: {e}")
        
        return charts
    
    def compile_report(self,
                      report_content: Dict,
                      charts: Dict,
                      output_format: str = 'pdf',
                      output_path: Optional[str] = None) -> Dict:
        """
        编译报告
        
        Args:
            report_content: 报告内容
            charts: 图表数据
            output_format: 输出格式
            output_path: 输出路径
            
        Returns:
            Dict: 报告生成结果
        """
        if output_format == 'pdf':
            return self.export_to_pdf(report_content, charts, output_path)
        elif output_format == 'html':
            return self.export_to_html(report_content, charts, output_path)
        elif output_format == 'markdown':
            return self.export_to_markdown(report_content, charts, output_path)
        elif output_format == 'docx':
            return self.export_to_docx(report_content, charts, output_path)
        else:
            # 默认使用PDF
            return self.export_to_pdf(report_content, charts, output_path)
    
    def export_to_pdf(self,
                     report_content: Dict,
                     charts: Dict,
                     output_path: Optional[str] = None) -> Dict:
        """
        导出为PDF
        
        Args:
            report_content: 报告内容
            charts: 图表数据
            output_path: 输出路径
            
        Returns:
            Dict: 导出结果
        """
        try:
            from matplotlib.backends.backend_pdf import PdfPages
            
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                symbol = report_content['metadata']['symbol']
                output_path = f"{symbol}_report_{timestamp}.pdf"
            
            with PdfPages(output_path) as pdf:
                # 封面页
                self.create_cover_page(pdf, report_content)
                
                # 目录页
                self.create_table_of_contents(pdf, report_content)
                
                # 内容页
                for section in report_content['sections']:
                    self.create_section_page(pdf, section, charts)
                
                # 图表页
                for chart_name, chart_data in charts.items():
                    if 'fig' in chart_data:
                        pdf.savefig(chart_data['fig'], bbox_inches='tight')
            
            return {
                'success': True,
                'output_path': output_path,
                'format': 'pdf',
                'pages': len(report_content['sections']) + len(charts) + 2
            }
            
        except Exception as e:
            print(f"导出PDF失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'format': 'pdf'
            }
    
    def export_to_html(self,
                      report_content: Dict,
                      charts: Dict,
                      output_path: Optional[str] = None) -> Dict:
        """
        导出为HTML
        
        Args:
            report_content: 报告内容
            charts: 图表数据
            output_path: 输出路径
            
        Returns:
            Dict: 导出结果
        """
        try:
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                symbol = report_content['metadata']['symbol']
                output_path = f"{symbol}_report_{timestamp}.html"
            
            # 创建HTML内容
            html_content = self.create_html_content(report_content, charts)
            
            # 保存HTML文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'output_path': output_path,
                'format': 'html'
            }
            
        except Exception as e:
            print(f"导出HTML失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'format': 'html'
            }
    
    def export_to_markdown(self,
                          report_content: Dict,
                          charts: Dict,
                          output_path: Optional[str] = None) -> Dict:
        """
        导出为Markdown
        
        Args:
            report_content: 报告内容
            charts: 图表数据
            output_path: 输出路径
            
        Returns:
            Dict: 导出结果
        """
        try:
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                symbol = report_content['metadata']['symbol']
                output_path = f"{symbol}_report_{timestamp}.md"
            
            # 创建Markdown内容
            md_content = self.create_markdown_content(report_content)
            
            # 保存Markdown文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            # 保存图表
            for chart_name, chart_data in charts.items():
                if 'fig' in chart_data:
                    chart_path = f"{output_path.replace('.md', '')}_{chart_name}.png"
                    chart_data['fig'].savefig(chart_path, dpi=self.dpi, bbox_inches='tight')
            
            return {
                'success': True,
                'output_path': output_path,
                'format': 'markdown'
            }
            
        except Exception as e:
            print(f"导出Markdown失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'format': 'markdown'
            }
    
    def export_to_docx(self,
                      report_content: Dict,
                      charts: Dict,
                      output_path: Optional[str] = None) -> Dict:
        """
        导出为DOCX
        
        Args:
            report_content: 报告内容
            charts: 图表数据
            output_path: 输出路径
            
        Returns:
            Dict: 导出结果
        """
        try:
            # 需要安装python-docx
            import docx
            from docx.shared import Inches
            
            if output_path is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                symbol = report_content['metadata']['symbol']
                output_path = f"{symbol}_report_{timestamp}.docx"
            
            # 创建文档
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading(report_content['title'], 0)
            doc.add_paragraph(report_content['subtitle'])
            
            # 添加元数据
            for key, value in report_content['metadata'].items():
                doc.add_paragraph(f"{key}: {value}")
            
            # 添加章节
            for section in report_content['sections']:
                doc.add_heading(section['title'], level=1)
                doc.add_paragraph(section['content'])
            
            # 添加图表
            for chart_name, chart_data in charts.items():
                if 'fig' in chart_data:
                    # 保存临时图片
                    temp_path = f"temp_{chart_name}.png"
                    chart_data['fig'].savefig(temp_path, dpi=self.dpi, bbox_inches='tight')
                    
                    # 添加到文档
                    doc.add_heading(f"图表: {chart_name}", level=2)
                    doc.add_picture(temp_path, width=Inches(6))
            
            # 保存文档
            doc.save(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'format': 'docx'
            }
            
        except ImportError:
            print("导出DOCX需要安装python-docx: pip install python-docx")
            return {
                'success': False,
                'error': 'python-docx not installed',
                'format': 'docx'
            }
        except Exception as e:
            print(f"导出DOCX失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'format': 'docx'
            }
    
    # 以下是一些辅助方法，用于生成报告内容
    
    def generate_executive_summary(self,
                                  data: pd.DataFrame,
                                  symbol: str,
                                  analysis_results: Dict) -> str:
        """生成执行摘要"""
        latest_price = data['Close'].iloc[-1]
        price_change = ((latest_price - data['Close'].iloc[0]) / data['Close'].iloc[0]) * 100
        
        summary = f"执行摘要: {symbol} 当前价格 ${latest_price:.2f}, 期间变化 {price_change:+.2f}%"
